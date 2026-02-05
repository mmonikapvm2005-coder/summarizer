from flask import Flask, render_template, request, send_file
from groq import Groq
from io import BytesIO
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
import PyPDF2
import os

app = Flask(__name__)

# ✅ Proper Groq client initialization
client = Groq(api_key=os.getenv("GROQ_API_KEY"))

# ---------- Helpers ----------

def extract_text(file):
    filename = file.filename.lower()

    if filename.endswith(".txt"):
        return file.read().decode("utf-8")

    elif filename.endswith(".pdf"):
        reader = PyPDF2.PdfReader(file)
        return "\n".join(page.extract_text() or "" for page in reader.pages)

    elif filename.endswith(".docx"):
        doc = Document(file)
        return "\n".join(p.text for p in doc.paragraphs)

    else:
        return ""

def summarize_text(text):
    response = client.chat.completions.create(
        model="llama-3.1-8b-instant",
        messages=[
            {"role": "system", "content": "Summarize the text clearly and concisely."},
            {"role": "user", "content": text}
        ],
        temperature=0.3
    )
    return response.choices[0].message.content.strip()

# ---------- Routes ----------

@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        file = request.files.get("file")
        format_type = request.form.get("format")

        if not file:
            return "No file uploaded", 400

        text = extract_text(file)

        if not text.strip():
            return "Could not extract text from file", 400

        summary = summarize_text(text)

        # -------- TXT --------
        if format_type == "txt":
            buffer = BytesIO(summary.encode("utf-8"))
            buffer.seek(0)
            return send_file(
                buffer,
                as_attachment=True,
                download_name="summary.txt",
                mimetype="text/plain"
            )

        # -------- PDF --------
        if format_type == "pdf":
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer)
            styles = getSampleStyleSheet()
            story = [Paragraph(line, styles["Normal"]) for line in summary.split("\n")]
            doc.build(story)
            buffer.seek(0)
            return send_file(
                buffer,
                as_attachment=True,
                download_name="summary.pdf",
                mimetype="application/pdf"
            )

        # -------- WORD --------
        if format_type == "word":
            buffer = BytesIO()
            doc = Document()
            for line in summary.split("\n"):
                doc.add_paragraph(line)
            doc.save(buffer)
            buffer.seek(0)
            return send_file(
                buffer,
                as_attachment=True,
                download_name="summary.docx",
                mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        return "Invalid format selected", 400

    return render_template("index.html")

if __name__ == "__main__":
    app.run(debug=True)
